from flask import Blueprint, jsonify, request, send_file
from sqlalchemy.orm import joinedload
from models import db, Checklist, ChecklistCompletion, ChecklistItem, ChecklistItemResponse, User, Project, ProjectUser
from auth import token_required
from datetime import datetime, date, timezone, timedelta
from docx import Document
from notification_service import create_notification
from risk_calculator import recalculate_project_risk
import io
import os
import json
from werkzeug.utils import secure_filename

checklist_bp = Blueprint('checklist', __name__)

UPLOAD_FOLDER = '/app/uploads/checklist_photos'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@checklist_bp.route('/api/checklists', methods=['GET'])
@token_required
def get_checklists():
    """Получение списка всех чек-листов"""
    try:
        checklists = Checklist.query.all()
        return jsonify([checklist.to_dict() for checklist in checklists]), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/checklists/pending-approval', methods=['GET'])
@token_required
def get_pending_approvals():
    """Получение чек-листов, ожидающих согласования (для инспектора)"""
    try:
        current_user_role = request.current_user.get('role')
        
        if current_user_role != 'inspector':
            return jsonify({'error': 'Только инспектор может просматривать чек-листы на согласование'}), 403
        
        pending_completions = ChecklistCompletion.query.filter_by(
            approval_status='pending'
        ).order_by(ChecklistCompletion.completion_date.desc()).all()
        
        result = []
        for completion in pending_completions:
            completion_dict = completion.to_dict()
            completion_dict['checklist'] = completion.checklist.to_dict()
            completion_dict['project'] = {
                'id': completion.project.id,
                'name': completion.project.name,
                'address': completion.project.address
            }
            completion_dict['completed_by'] = {
                'id': completion.completed_by.id,
                'first_name': completion.completed_by.first_name,
                'last_name': completion.completed_by.last_name,
                'email': completion.completed_by.email
            }
            completion_dict['item_responses'] = [resp.to_dict() for resp in completion.item_responses]
            result.append(completion_dict)
        
        return jsonify(result), 200
        
    except Exception as e:
        return jsonify({'error': str(e), 'message': 'Ошибка получения чек-листа'}), 500


@checklist_bp.route('/api/checklists/<int:checklist_id>', methods=['GET'])
@token_required
def get_checklist(checklist_id):
    """Получение детальной информации о чек-листе"""
    try:
        checklist = Checklist.query.get_or_404(checklist_id)
        return jsonify(checklist.to_dict()), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 404


@checklist_bp.route('/api/checklists/<int:checklist_id>/complete', methods=['POST'])
@token_required
def complete_checklist(checklist_id):
    """Заполнение чек-листа с загрузкой файлов."""
    try:
        current_user_role = request.current_user.get('role')
        current_user_id = request.current_user.get('id')
        
        if current_user_role != 'client':
            return jsonify({'error': 'Только клиент может заполнять чек-листы'}), 403

        checklist = Checklist.query.get_or_404(checklist_id)
        project_id = request.form.get('project_id', type=int)
        if not project_id:
            return jsonify({'error': 'project_id is required'}), 400

        if checklist.type == 'daily':
            last_completion = ChecklistCompletion.query.filter_by(
                project_id=project_id,
                checklist_id=checklist_id
            ).order_by(ChecklistCompletion.completion_date.desc()).first()

            if last_completion:
                completion_time = last_completion.completion_date.replace(tzinfo=timezone.utc)
                if (datetime.now(timezone.utc) - completion_time) < timedelta(hours=12):
                    return jsonify({'error': 'Ежедневный чек-лист можно создавать не чаще, чем раз в 12 часов'}), 429

            project = Project.query.get(project_id)
            if not project:
                return jsonify({'error': 'Проект не найден'}), 404
            if project.status != 'active':
                return jsonify({'error': 'Объект должен быть активирован перед заполнением ежедневного чек-листа'}), 403

        photo_paths = []
        uploaded_files = request.files.getlist('photos')
        for file in uploaded_files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                random_chars = os.urandom(4).hex()
                unique_filename = f"{timestamp}_{random_chars}_{filename}"
                file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
                file.save(file_path)
                photo_paths.append(f'checklist_photos/{unique_filename}')

        completion = ChecklistCompletion(
            checklist_id=checklist_id,
            project_id=project_id,
            completed_by_id=current_user_id,
            completion_date=datetime.now(timezone.utc),
            items_data=json.loads(request.form.get('items_data', '{}')),
            photos=photo_paths,
            geolocation=request.form.get('geolocation'),
            notes=request.form.get('notes'),
            approval_status='not_required' if not checklist.requires_approval else 'pending',
            initialization_required=checklist.requires_initialization,
            initialized_at=datetime.now(timezone.utc) if checklist.requires_initialization else None,
            initialization_geolocation=request.form.get('initialization_geolocation') if checklist.requires_initialization else None
        )
        db.session.add(completion)
        db.session.flush()

        responses_data = json.loads(request.form.get('responses', '[]'))
        for resp in responses_data:
            response = ChecklistItemResponse(
                completion_id=completion.id,
                item_id=resp['item_id'],
                response=resp['response'],
                photos=resp.get('photos', []),
                comment=resp.get('comment'),
                geolocation=resp.get('geolocation', request.form.get('geolocation')),
                created_at=datetime.now(timezone.utc)
            )
            db.session.add(response)

        db.session.commit()

        recalculate_project_risk(project_id, triggering_user_id=request.current_user['id'])
        
        if checklist.requires_approval and checklist.type == 'opening':
            project = db.session.get(Project, project_id)
            inspector_users = db.session.query(User).filter_by(role='inspector', is_active=True).all()
            
            for inspector in inspector_users:
                create_notification(
                    user_id=inspector.id,
                    message=f"Новый чек-лист открытия объекта '{project.name}' ожидает согласования",
                    link=f"/pending-approvals"
                )
        
        return jsonify(completion.to_dict()), 201

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/projects/<int:project_id>/checklists/<checklist_type>/completions', methods=['GET'])
@token_required
def get_completions(project_id, checklist_type):
    """Получение истории заполнений чек-листов для проекта"""
    try:
        checklist = Checklist.query.filter_by(type=checklist_type).first()
        if not checklist:
            return jsonify({'error': 'Чек-лист не найден'}), 404

        completions = ChecklistCompletion.query.filter_by(
            project_id=project_id,
            checklist_id=checklist.id
        ).order_by(ChecklistCompletion.completion_date.desc()).all()

        result = []
        for completion in completions:
            completion_dict = completion.to_dict()
            completion_dict['checklist'] = completion.checklist.to_dict()
            completion_dict['completed_by'] = {
                'id': completion.completed_by.id,
                'first_name': completion.completed_by.first_name,
                'last_name': completion.completed_by.last_name,
                'email': completion.completed_by.email
            }
            if completion.approved_by:
                completion_dict['approved_by'] = {
                    'id': completion.approved_by.id,
                    'first_name': completion.approved_by.first_name,
                    'last_name': completion.approved_by.last_name
                }
            result.append(completion_dict)

        return jsonify(result), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/checklist-completions/<int:completion_id>/details', methods=['GET'])
@token_required
def get_completion_details(completion_id):
    """Получение детальной информации о заполненном чек-листе"""
    try:
        completion = ChecklistCompletion.query.get_or_404(completion_id)
        
        completion_dict = completion.to_dict()
        completion_dict['checklist'] = completion.checklist.to_dict()
        completion_dict['project'] = {
            'id': completion.project.id,
            'name': completion.project.name,
            'address': completion.project.address
        }
        completion_dict['completed_by'] = {
            'id': completion.completed_by.id,
            'first_name': completion.completed_by.first_name,
            'last_name': completion.completed_by.last_name,
            'email': completion.completed_by.email
        }
        if completion.approved_by:
            completion_dict['approved_by'] = {
                'id': completion.approved_by.id,
                'first_name': completion.approved_by.first_name,
                'last_name': completion.approved_by.last_name
            }

        completion_dict['item_responses'] = [resp.to_dict() for resp in completion.item_responses]

        return jsonify(completion_dict), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 404


@checklist_bp.route('/api/checklist-completions/<int:completion_id>/export', methods=['GET'])
@token_required
def export_checklist_docx(completion_id):
    """Экспорт заполненного чек-листа в DOCX"""
    try:
        completion = ChecklistCompletion.query.get_or_404(completion_id)
        checklist = completion.checklist

        if checklist.type == 'opening':
            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'checklist_opening_template.docx')
        else:
            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'checklist_daily_template.docx')

        if not os.path.exists(template_path):
            return jsonify({'error': f'Шаблон не найден: {template_path}'}), 404

        doc = Document(template_path)

        completion_date_str = completion.completion_date.strftime('%d.%m.%Y') if completion.completion_date else datetime.now().strftime('%d.%m.%Y')
        completed_by_name = f"{completion.completed_by.last_name} {completion.completed_by.first_name}"
        
        foreman_name = ''
        if completion.project:
            foreman_member = next((m for m in completion.project.members if m.user and m.user.role == 'foreman'), None)
            if foreman_member and foreman_member.user:
                foreman_name = f"{foreman_member.user.last_name} {foreman_member.user.first_name}"
        
        replacements = {
            '{{PROJECT_NAME}}': completion.project.name if completion.project else '',
            '{{DATE}}': completion_date_str,
            '{{COMPLETED_BY}}': completed_by_name,
            '{{ADDRESS}}': completion.project.address if completion.project else '',
            '{{NOTES}}': completion.notes or '',
            '{{FOREMAN}}': foreman_name,
            '{{CHECK_DATE}}': completion_date_str
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

        responses = {r.item_id: r for r in completion.item_responses}
        sorted_items = sorted(checklist.items, key=lambda x: x.order)

        if len(doc.tables) > 0:
            table = doc.tables[0]
            
            for idx, item in enumerate(sorted_items):
                row_idx = idx + 2
                
                while row_idx >= len(table.rows):
                    table.add_row()
                
                row = table.rows[row_idx]
                
                try:
                    row.cells[0].text = str(idx + 1)
                except:
                    pass
                
                try:
                    row.cells[1].text = item.text
                except:
                    pass
                
                response = responses.get(item.id)
                
                try:
                    row.cells[2].text = '✓' if response and response.response == 'yes' else ''
                except:
                    pass
                
                try:
                    row.cells[3].text = '✓' if response and response.response == 'no' else ''
                except:
                    pass
                
                try:
                    row.cells[4].text = '✓' if response and response.response == 'not_applicable' else ''
                except:
                    pass
                
                try:
                    row.cells[5].text = response.comment if response and response.comment else ''
                except:
                    pass
        
        doc.add_paragraph()
        doc.add_paragraph("Проверил: _____________/_______________/")
        doc.add_paragraph()
        doc.add_paragraph("Чек-лист по выявленным нарушениям получил: _______________/_____________/")

        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)

        filename = f"checklist_{completion.id}_{completion_date_str.replace('.', '')}.docx"

        return send_file(
            docx_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/projects/<int:project_id>/checklists/<int:checklist_id>/today', methods=['GET'])
@token_required
def check_today_completion(project_id, checklist_id):
    """Проверка заполнения чек-листа сегодня"""
    try:
        today = date.today()
        completion = ChecklistCompletion.query.filter(
            ChecklistCompletion.project_id == project_id,
            ChecklistCompletion.checklist_id == checklist_id,
            db.func.date(ChecklistCompletion.completion_date) == today
        ).first()

        if completion:
            return jsonify({
                'completed': True,
                'completion': completion.to_dict()
            }), 200
        else:
            return jsonify({'completed': False}), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/projects/<int:project_id>/checklists/<int:checklist_id>/initialize', methods=['POST'])
@token_required
def initialize_checklist(project_id, checklist_id):
    """Инициализация чек-листа с фиксацией времени и геолокации начала заполнения"""
    try:
        current_user_id = request.current_user.get('id')
        current_user_role = request.current_user.get('role')
        
        if current_user_role != 'client':
            return jsonify({'error': 'Только клиент может инициализировать чек-листы'}), 403
        
        checklist = Checklist.query.get_or_404(checklist_id)
        data = request.get_json()
        geolocation = data.get('geolocation')
        
        today = date.today()
        existing_completion = ChecklistCompletion.query.filter(
            ChecklistCompletion.project_id == project_id,
            ChecklistCompletion.checklist_id == checklist_id,
            db.func.date(ChecklistCompletion.completion_date) == today
        ).first()
        
        if existing_completion:
            return jsonify({
                'success': True,
                'completion': existing_completion.to_dict(),
                'message': 'Чек-лист уже инициализирован сегодня'
            }), 200
        
        completion = ChecklistCompletion(
            checklist_id=checklist_id,
            project_id=project_id,
            completed_by_id=current_user_id,
            completion_date=datetime.now(timezone.utc),
            items_data={},
            geolocation=geolocation,
            initialization_required=checklist.requires_initialization,
            initialized_at=datetime.now(timezone.utc),
            initialization_geolocation=geolocation,
            approval_status='not_required' if not checklist.requires_approval else 'pending'
        )
        
        db.session.add(completion)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'completion': completion.to_dict()
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/checklist-completions/<int:completion_id>/update', methods=['PUT'])
@token_required
def update_completion(completion_id):
    """Редактирование заполненного чек-листа"""
    try:
        current_user_id = request.current_user.get('id')
        completion = ChecklistCompletion.query.get_or_404(completion_id)

        if completion.completed_by_id != current_user_id:
            return jsonify({'error': 'Недостаточно прав'}), 403
        
        if completion.checklist.type == 'daily':
            project = completion.project
            if not project or project.status != 'active':
                return jsonify({'error': 'Объект должен быть активирован для редактирования ежедневного чек-листа'}), 403

        data = request.get_json()
        
        if 'notes' in data:
            completion.notes = data['notes']
        if 'photos' in data:
            completion.photos = data['photos']

        if 'responses' in data:
            ChecklistItemResponse.query.filter_by(completion_id=completion_id).delete()
            
            for resp in data['responses']:
                response = ChecklistItemResponse(
                    completion_id=completion.id,
                    item_id=resp['item_id'],
                    response=resp['response'],
                    photos=resp.get('photos', []),
                    comment=resp.get('comment'),
                    geolocation=resp.get('geolocation'),
                    created_at=datetime.now(timezone.utc)
                )
                db.session.add(response)

        db.session.commit()
        return jsonify(completion.to_dict()), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/checklist-completions/<int:completion_id>/approve', methods=['POST'])
@token_required
def approve_completion(completion_id):
    """Согласование чек-листа (только для инспектора)"""
    try:
        current_user_role = request.current_user.get('role')
        current_user_id = request.current_user.get('id')
        
        if current_user_role != 'inspector':
            return jsonify({'error': 'Только инспектор может согласовывать чек-листы'}), 403

        completion = ChecklistCompletion.query.options(
            joinedload(ChecklistCompletion.checklist)
        ).get_or_404(completion_id)
        
        completion.approval_status = 'approved'
        completion.approved_by_id = current_user_id
        completion.approved_at = datetime.now(timezone.utc)
        
        if completion.checklist.type == 'opening':
            project = completion.project
            project.status = 'active'
        
        db.session.commit()

        recalculate_project_risk(completion.project_id, triggering_user_id=request.current_user['id'])
        
        project = completion.project
        create_notification(
            user_id=completion.completed_by_id,
            message=f"Чек-лист для проекта '{project.name}' согласован",
            link=f"/projects/{project.id}"
        )
        
        return jsonify(completion.to_dict()), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@checklist_bp.route('/api/checklist-completions/<int:completion_id>/reject', methods=['POST'])
@token_required
def reject_completion(completion_id):
    """Отклонение чек-листа (только для инспектора)"""
    try:
        current_user_role = request.current_user.get('role')
        current_user_id = request.current_user.get('id')
        
        if current_user_role != 'inspector':
            return jsonify({'error': 'Только инспектор может отклонять чек-листы'}), 403

        completion = ChecklistCompletion.query.options(
            joinedload(ChecklistCompletion.checklist)
        ).get_or_404(completion_id)
        data = request.get_json()
        
        completion.approval_status = 'rejected'
        completion.approved_by_id = current_user_id
        completion.approved_at = datetime.now(timezone.utc)
        completion.rejection_reason = data.get('reason', '')
        
        db.session.commit()

        recalculate_project_risk(completion.project_id, triggering_user_id=request.current_user['id'])
        
        project = completion.project
        create_notification(
            user_id=completion.completed_by_id,
            message=f"Чек-лист для проекта '{project.name}' отклонен",
            link=f"/projects/{project.id}"
        )
        
        return jsonify(completion.to_dict()), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500
